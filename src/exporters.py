import re
from io import BytesIO

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def generate_pdf(text):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=60,
        bottomMargin=40,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CustomBold",
            fontSize=12,
            leading=16,
            fontName="Helvetica-Bold",
            textColor="#222244",
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CustomBody",
            fontSize=10,
            leading=14,
            fontName="Helvetica",
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CustomBullet",
            fontSize=10,
            leading=14,
            leftIndent=12,
            bulletIndent=0,
            bulletFontName="Helvetica-Bold",
        )
    )

    flowables = []

    for raw_line in text.split("\n"):
        line = raw_line.strip().replace("**", "")

        if line.startswith("# "):
            flowables.append(Paragraph(line.replace("# ", ""), styles["CustomBold"]))
            flowables.append(Spacer(1, 10))
        elif line.startswith("## "):
            flowables.append(Paragraph(line.replace("## ", ""), styles["CustomBold"]))
            flowables.append(Spacer(1, 8))
        elif re.match(r"^\d+\.", line):
            flowables.append(Paragraph(line, styles["CustomBold"]))
            flowables.append(Spacer(1, 8))
        elif line.startswith("- "):
            flowables.append(Paragraph(f"• {line[2:]}", styles["CustomBullet"]))
            flowables.append(Spacer(1, 4))
        elif not line:
            flowables.append(Spacer(1, 10))
        else:
            flowables.append(Paragraph(line, styles["CustomBody"]))
            flowables.append(Spacer(1, 6))

    doc.build(flowables)
    buffer.seek(0)
    return buffer


def generate_docx(text):
    doc = Document()
    doc.add_heading("GitHub Portfolio Suggestions", level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_txt(text):
    return BytesIO(text.encode("utf-8"))


def generate_markdown(text):
    return BytesIO(text.encode("utf-8"))
